"""DOCX generation service using python-docx directly.

Generates a structured Word document from the parsed XML context.
No template file needed — full programmatic control.
"""

from __future__ import annotations

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BRAND = RGBColor(0x5D, 0x5E, 0x60)
DARK = RGBColor(0x3A, 0x3B, 0x3D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)


def _shade_cell(cell, hex_color: str = "5D5E60"):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"})
    tc_pr.append(shd)


def _add_branded_heading(doc: Document, text: str, level: int = 1, color: RGBColor = BRAND):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _add_key_value_table(doc: Document, rows: list[dict]):
    """Add a 2-column table (Campo | Valor) with branded header."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0]
    for i, label in enumerate(["Campo", "Valor"]):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        _shade_cell(cell)

    # Data rows
    for row_data in rows:
        campo = str(row_data.get("campo", ""))
        valor = str(row_data.get("valor", ""))
        if campo == "---":
            continue
        row = table.add_row()
        row.cells[0].text = campo
        row.cells[1].text = valor
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")


def _render_sections(doc: Document, sections: list[dict], level: int = 2):
    """Recursively render sections as headings + tables."""
    for section in sections:
        title = section.get("title", "Sin título")
        rows = section.get("rows", [])
        children = section.get("children", [])

        _add_branded_heading(doc, title, level=level, color=DARK if level == 2 else BRAND)

        if rows:
            _add_key_value_table(doc, rows)

        if children:
            _render_sections(doc, children, level=min(level + 1, 4))


def render_docx(context: dict, template_code: str = "ETL_Template") -> bytes:
    """Generate a complete DOCX document from the parsed XML context."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    project_name = context.get("project_name", "Sin nombre")
    root_tag = context.get("root_tag", "documento")
    sections = context.get("sections", [])
    flat = context.get("flat", {})

    # === Cover page ===
    for _ in range(3):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documentador")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Reporte generado automáticamente")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyecto: ")
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(project_name)
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Origen XML: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(f"<{root_tag}>")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Plantilla: {template_code}")
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND

    doc.add_page_break()

    # === Sections (structured) ===
    if sections:
        _add_branded_heading(doc, "Contenido estructurado", level=1)
        doc.add_paragraph("Secciones encontradas en el XML procesado:")
        doc.add_paragraph("")
        _render_sections(doc, sections)

    # === Flat key-value (all leaf nodes) ===
    if flat:
        doc.add_page_break()
        _add_branded_heading(doc, "Resumen de campos (vista plana)", level=1)
        doc.add_paragraph("Todos los campos de hoja del XML con su ruta completa:")
        doc.add_paragraph("")

        flat_rows = [{"campo": k, "valor": v} for k, v in flat.items()]
        _add_key_value_table(doc, flat_rows)

    # === Footer ===
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generado por Documentador — Transformación real de XML a DOCX")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND
    run.italic = True

    # Serialize
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
