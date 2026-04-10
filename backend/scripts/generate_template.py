"""Generate the ETL_Template.docx for docxtpl with Jinja2 tags."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BRAND = RGBColor(0x5D, 0x5E, 0x60)
DARK = RGBColor(0x3A, 0x3B, 0x3D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, hex_color="5D5E60"):
    from docx.oxml.ns import qn
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"})
    tc_pr.append(shd)


def make_template():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    # --- Title page ---
    for _ in range(3):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documentador")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Reporte generado automáticamente")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyecto: ")
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run("{{ project_name }}")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Origen XML: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("{{ root_tag }}")
    run.font.size = Pt(11)

    doc.add_page_break()

    # --- Content heading ---
    h = doc.add_heading("Contenido del documento", level=1)
    for run in h.runs:
        run.font.color.rgb = BRAND

    doc.add_paragraph("Secciones encontradas en el XML procesado:")
    doc.add_paragraph("")

    # --- Flat key-value table ---
    # Use a simple 2-col table with row loop via {%tr %}
    # {%tr %} MUST be the first thing in a table cell for docxtpl
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0]
    for i, text in enumerate(["Campo", "Valor"]):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        _shade_cell(cell)

    # Data row with {%tr %} tag — must be first content in the cell
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[0].paragraphs[0].add_run("{%tr for key, val in flat.items() %}{{ key }}")
    row.cells[1].text = ""
    row.cells[1].paragraphs[0].add_run("{{ val }}{%tr endfor %}")

    doc.add_paragraph("")

    # --- Footer ---
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generado por Documentador — Transformación real de XML a DOCX")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND
    run.italic = True

    out_dir = os.path.join(os.path.dirname(__file__), "..", "app", "templates")
    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, "ETL_Template.docx")
    doc.save(path)
    print(f"Template saved to {path}")


if __name__ == "__main__":
    make_template()
